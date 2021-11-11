import docx
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
# from .texts import x,y,xx
x1="______________________	                                                          Тошкент шахар."
x = """
Ш А Р Т Н О М А № 
Тураруй-жой биноси қурилишига инвестицион бадал тўғрисидаги
"""
y = """
        Кейинги матнда “Жамият” деб юритилувчи юридик шахс “INSOFT INVEST” масъулияти чекланган жамияти (СТИР: 304757296), номидан устав асосида иш олиб борадиган жамият Директори М.А.Очилов, бир тарафдан, ҳамда кейинги матнда “Инвестор” деб юритилувчи жисмоний шахс, Ўзбекистон Республикасининг фуқароси (_________________________________) _______________________________________ (СТИР: _____________ ЖШШИР: ___________________) __ __________ серия раками ______________________________ томонидан _______________ йил берилган рақамли паспорт эгаси, иккинчи тарафдан, биргаликда “Тарафлар” деб номланиб, Ўзбекистон Республикаси юритиб, келгуси матнда “Шартнома” деб юритилувчи ушбу Турар-жой биноси қурилишига инвестицион бадал тўғрисидаги  шартномани қуйидагилар тўғрисида туздилар
"""
xx="""
1. Шартнома предмети
"""
yy="""
        1.1. . Жамият ўз кучи ва (ёки) бошқа шахсларни – пудратчи ташкилотларни жалб қилган ҳолда кўп қаватли турар-жой бинолари мажмуасини (кейинги ўринларда – “Турар жой биноси”), Тошкент шаҳар, Яккасарой тумани, Жамшид Шоший 1 тор кўчасида жойлашган, Тошкент шаҳар ҳокимининг 07.03.2018 йилдаги 388-сонли қарори билан қурилиш учун ажратилган ер майдонида қуриш мажбуриятини ўз зиммасига олади. 
        Тасдиқланган қурилиш режасига мувофиқ  замонавий, кўп қаватли турар жой биносининг қурилиши, пастки қаватларида офис ва ижтимоий-маиший объектларнинг жойлашиши билан амалга оширилади. 
        Турар жой биноси давлат рўйхатидан ўтказилиб, балансга қабул қилинганидан ва фойдаланишга топширилганидан сўнг, Жамият Инвесторга хонадонни (Объект) топшириш мажбуриятини олади. 
        Инвестор, ушбу шартнома билан ўрнатилган муддат ва тартибда қурилишга инвестиция киритиш ва инвестиция Объектини ўз эгалигига қабул қилиш мажбуриятини олади.
        1.2. Қурилиш бўйича ишларни амалга ошириш ва молиялаштириш жадвали билан ўрнатилган тартибда, Объектни топшириш муддати, тарафлар томонидан тахминан __ойдан ____ ойигача деб белгиланмоқда.Ушбу мажбуриятни бажаришга тўсқинлик қилувчи, Жамиятнинг хоҳиш истагига боғлиқ бўлмаган вазиятлар вужудга келганда, шу жумладан, Давлат қабул комиссиясининг далолатномаси имзоланиш санасига боғлиқ ҳолда, Объектни топшириш муддати ўзгартирилиши мумкин. Ушбу шартномага мувофиқ, Турар жой биноси фойдаланишга топширилганидан сўнг Инвесторга ажратилиши лозим бўлган Объект, шу жумладан, Инвесторнинг маблағларини жалб қилган равишда қурилаётган турар жой биноси таркибига кирувчи турар жой хонадонидан иборат.
        1.3. Тарафлар, ушбу шартномага асосан Жамият Инвесторга топшириш-қабул қилиш далолатномаси асосида қуйидаги манзилда жойлашган хонадонни топшириши шартлиги ҳақида келишиб олдилар: Тошкент шаҳри, Яккасарой тумани, Аския №______ объектидаги, Жамшид Шоший _ тор кўчасида қурилаётган турар-жой биносининг, ___Блок ____қавати, ______ кв.метр майдонга эга  бўлган  _____-хонадон.	
        1.4. Жамият, ушбу шартномани имзолаш вақтига 1.3.-бандда кўрсатилган инвестиция Объекти-суд тортишувининг предмети эмаслигига; таъқиқ остида эмаслигига; гаров предмети эмаслигига; учинчи шахслар ҳуқуқи ва талаблари билан юкланмаганлигигаҳамда инвестиция Объектига учинчи шахслар ҳуқуқини вужудга келтириши мумкин бўлган бошқа Инвестиция бадали тўғрисидаги шартнома тузилмаганлигига кафолат беради.
        1.5. Тарафлар, шартноманинг 1.3. бандида кўрсатилган инвестиция Объектининг майдони, шу жумладан, унда жойлашган алоҳида хоналарнинг майдони лойиҳавий (тахминий) эканлигига ва Турар жой биноси қурилиши якунига етганидан сўнг кўпайиш ёки камайиш тарафига ўзгариши мумкинлигига келишдилар. Объектнинг (шу жумладан, ундаги алоҳида хоналарнинг) якуний майдони Тошкент шаҳар ер тузиш ва кўчмас мулк Кадастрининг Бош бошқармасининг ўлчовлари асосида белгиланади.   
        1.6. Ушбу Шартноманинг 1.3. бандида кўрсатилган инвестиция Объектининг майдони квадратураси ҳажми, Тошкент шаҳар ер тузиш ва кўчмас мулк Кадастрининг Бош бошқармасининг ўлчов маълумотларига нисбатан ошганда ёки камайганда, инвестицион бадалининг умумий қиймати қайта ҳисобланади ва ушбу шартномага асосан инвестиция маблағларининг суммаси ўзгаради.
"""
yyy="""
    2. Инвестиция Объектининг тавсифи, ишлар таркиби ва ҳажми
"""
yyyy="""
        2.1. Инвестор эгалигига инвестицион бадалдан фойдаланиш муддати якунига етганидан сўнг топшириладиган инвестиция Объекти бўйича ишлар ҳажми ва таркиби қуйидаги асосий шартларни инобатга олган ҳолда лойиҳа-смета ҳужжатларига мувофиқ келиши лозим: 
        2.1.1. Штукатурка, бўёқ ишлари, хоналараро тўсиқлар ва эшикларни ўрнатиш;  ҳамда иккинчи босқичнинг махсус монтаж ишлари: унитаз, ванна, раковина, сув смесителлари, квартира ичидаги электр ишлари ва розеткалар, чироқни ёқиш-ўчириш мосламалари, ёритиш ва ўлчаш ускуналарини ўрнатиш Инвестор томонидан мустақил амалга оширилади ва умумий инвестицион бадал (шартнома суммаси) ичига кирмайди;
        2.1.2. Жамият Объектга нисбатан фақат қуйидаги ишларни амалга оширади: 
        - ташқи деворлар бўйлаб дераза блокларини ва киришда темир эшикни ўрнатиш (ички хоналараро эшиклар ва ички дераза блокларини ўрнатиш ушбу шартнома бўйича ишлар таркибига кирмайди);
        - электрмонтаж ишлари: зинапоя майдонида жойлашган электр қутисига электр сими узатилиб, хонадон ичида электр симларининг монтажи амалга оширилмайди, розетка ва ёқиш-ўчириш мосламалари ўрнатилмайди;  
        -  сантехника ишлари: совуқ сув канализация қувурлар орқали ўтказилади, сантехника воситалари (раковина, ванна, унитаз, сув кранлари, сув ўлчаш асбоблари ва ҳакозолар) ўрнатилмайди;
        - газлаштириш: газ плиталари ва газ ўлчаш асбобларини ўрнатмаган ҳолда, маиший газ ошхонага қувур орқали ўтказилади;
        - телефонлаштириш: телефон сими зинапоя майдончасида жойлашган қутига ўтказилади;
        2.1.3. Тарафлар, Объект қурилиши бўйича барча бирламчи рухсатномавий, ҳуқуқни ўрнатувчи, лойиҳа-смета ҳужжатларни тўлиқ ҳажмда, Жамиятнинг турар жой биносини қуришга бўлган ҳуқуқига ишонч ҳосил қилишга, Объектни аниқлаб олишга, унинг ўлчамлари ва майдонда мўлжал олишга, Объектнинг ҳам умумий турар жой биносининг ҳам асосий режавий, тузилмавий, функционал, меъморий-бадиий, технологик, санитария-гигиена ва муҳандислик кўрсаткичларини аниқлаб олишга; инвестиция Объектининг умумий ва нотурар жойларининг таркиби ва нима учун мўлжалланганлигини белгилаб олишга етарли бўлган маълумотлар Жамият томонидан тақдим қилинганлигини ва Инвестор томонидан қабул қилинганлигини тасдиқлайдилар. 
        2.1.4. Инвестор, юқорида кўрсатилган ахборотлар ҳажми у учун етарли эканлигини, ушбу шартномани тузиш қарорини қабул қилиш ва ўзининг инвестицияда иштирокидаги таваккалчилик даражасини аниқлаш учун асос сифатида хизмат қилганлигини ҳамда унинг ушбу шартномада кўрсатилган мажбуриятларини керакли тарзда бажариши учун етарли бўлган ҳозирги ва келгусидаги молиявий ҳолатини баҳолаб олганини тасдиқлайди. 
        2.1.5. Инвестор томонидан турар жой учун ҳуқуқни белгиловчи ҳужжатлар (Ўзбекистон Республикаси қонунчилигига мувофиқ кадастр ва бошқа ҳужжатлар) олинишига қадар, у инвестиция объектини қайта режалаштириш ва таъмирлаш ҳуқуқига эга бўлмайди.
        2.1.6. Инвестор, Ўзбекистон Республикасининг ҳуқуқий-меъёрий ҳужжатлари ва техник, санитария ва архитектура хавфсизлиги талабларига жавоб берадиган иккинчи босқич безак ва махсус монтаж ишларини, бинога зарар етказмаган ҳолда, ўз маблағлари ҳисобидан амалга ошириш, қонунчиликда ўрнатилган тартибда ва муддатларда коммунал (эксплуатация) хизматини кўрсатувчи хизматлар ва объектга хизмат кўрсатувчи бошқа ташкилотлар билан шартнома тузиш мажбуриятини олади.
        2.1.7. Агар Инвестор топшириғи ва Жамият розилиги билан иккинчи босқич ишларини амалга ошириш иккинчи Тарафга топширилса, Тарафлар ушбу шартномага ҳавола қилмаган тарзда янги шартнома тузадилар.
"""
xxx="""
    3. Тарафларнинг ҳуқуқ ва мажбуриятлари
"""
a="""
        3.1. Инвесторнинг ҳуқуқ ва мажбуриятлари:
        3.1.1. Инвестор 18 ёшдан ошган, муомалага лаёқатли (шартнома тузилиши вақтига ва тўлиқ ижро этилишига қадар) шахс бўлиши лозим;
        - ҳужжатларнинг асл нусхаларини (паспорт, СТИР, рўйхатга олинган жойи ва турар жойидан, барча оила аъзоларини кўрсатган ҳолда маълумотнома, бир йиллик иш ҳақи тўғрисида маълумотнома, меҳнат дафтарчаси) тақдим қилиш;
        - охирги иш жойида камида олти ойлик иш стажи;
        - Шартнома бўйича асосий қарз учун тўловни амалга ошириш имкониятини берувчи барқарор даромад (иш ҳақи, якка тартибдаги фаолиятдан олинган даромад, дивидендлар ва ҳакозо);
        - ушбу шартноманинг шартларини бажариш ҳақида Жамият раҳбарига ёзма равишда кафолати хати тақдим қилиш;
        - Жамият талаби билан, шартнома шартларини бажариш учун зарур бўлган бошқа ҳужжатларни тақдим этиш;
        - камида учта кафил кафилликни ёзма ариза шаклида нотиариал тасдиқланган ҳолда тақдим қилиш;
        Инвестор, ушбу шартномада кўрсатилган муддат ичида қурилишга ўз инвестицион бадалини тўлиқ киритиши шарт. 
        3.1.2. Инвестор, қурилишга инвестиция тўловларини банк орқали ўзига қулай усулда тўлаши ёки Жамият билан келишилган номенклатура ва миқдорда қурилиш моллари кўринишида амалга ошириши мумкин.
        3.1.3. Инвестор, Жамиятдан ушбу шартномага мувофиқ қурилиш якунига етганлиги ва Объект топширилиш учун тайёр эканлиги ҳақида хабарномани олган вақтидан бошлаб 10 (ўн) кун ичида Объектни топшириш-қабул қилиш далолатномаси асосида қабул қилиб олиши шарт. 
        Агар Инвестор ҳужжатлар билан тасдиқланган узрли сабаб кўрсатмай туриб, қабул қилиш муддатларини бузса, Жамиятнинг Объектни Инвесторга топшириш бўйича мажбурияти кўрсатилган муддат тугагандан кейин бажарилган деб ҳисобланади ва ушбу муддат ўтганидан сўнг, Объектнинг тасодифий шикастланиши ёки нобуд бўлиши таваккалчилиги Инвесторга ўтади.
        3.1.4. Инвестор топшириш-қабул қилиш далолатномасига имзо қўйган вақтидан бошлаб Объектга эгалик ҳуқуқини расмийлаштириш, шу жумладан, кўчмас мулкка бўлган ҳуқуқни давлат рўйхатидан ўтказиш ваколатига эга органлардаги ҳаражатларни ўз зиммасига олади ва қоплайди. 
        3.1.5. Инвестор квартиралар, турар жой бинолари ва бино олди ҳудудлардан фойдаланишнинг техник (реконструкция қилмаслик, таянч деворларини бузмаслик, ўтиш йўлларини очмаслик), ёнғин хавфсизлигига ва санитарияга оид қоида ва меъёрларга риоя қилиши, ускуналар (кондиционерлар, турли кўринишдаги антенналар, панжаралар) ўрнатганда уйнинг фасад қисми, умумий фойдаланиш майдони ва мулкининг кўринишини ўзгартирмаслиги шарт. 
        Ушбу ишларни амалга ошириш учун Жамият розилиги олинган, ёки бундай ишлар Инвесторнинг топшириғи билан Жамият томонидан амалга оширилган ҳоллар бундан мустасно.
        Акс ҳолда, қайта режалаштириш ва қайта тузишнинг кўринишидан қатъи назар, Инвесторнинг ҳаракатлари Жамиятга зарар етказган ҳаракатлар деб тан олиниб, Жамият Инвестор томонидан киритилган бадал ҳисобидан етказилган зарар суммасини ушлаб қолиш ҳуқуқига эга бўлади.
        Бунда, Инвестор қўшимча равишда, ушбу ҳолат билан боғлиқ барча салбий оқибатларни, шу жумладан, Жамиятнинг қайта режалаштириш ва қайта тузиш оқибатларини бартараф этиш билан боғлиқ ҳаражатларини мустақил равишда зиммасига олади ва қоплайди.
        3.1.6. Объектга хизмат кўрсатиш ва фойдаланиш ҳаражатларни қоплаш учун зарур бўлган маблағларни ўз вақтида амалга ошириб бориш. 
        3.1.7. Инвестор Жамиятни ёзма равишда ўз маълумотлари (паспорт маълумотлари, манзил, ҳисоб рақами, телефон рақамлари ва ҳакозо) ўзгарганлиги ҳақида рўй берган ўзгаришлардан кейин беш кун ичида хабардор қилиши лозим. 
        3.1.8. Инвестор, мустақил равишда ва қонунчилик, меъёрий ҳужжатлар билан ўрнатилган муддатда, қарздорликларга йўл қўймасдан Объект бўйича барча хизмат турлари учун тўлов киритиши лозим (шу жумладан, коммунал ва бошқа тўловлар).
        Жамият ушбу борада Инвесторнинг ҳар қандай қарздорлиги учун масъулиятни ўз зиммасига олмайди.
        3.1.9. Инвестор ушбу шартноманинг ўзининг инвестицион мажбуриятлари тўлиқ бажарилишга қадар Объектни ижарага (суб-ижарага) беришга ҳақли эмас.
        3.1.10. Инвестор ўз инвестицион мажбуриятларини муддатдан олдин сўндириши, ҳамда олдиндан тўлов киритиши мумкин.
        3.1.11. Агар Инвестор шартнома бўйича тўловларни амалга ошириш имкониятини йўқотса (шу жумладан, лаёқатини йўқотса, вафот этса) унинг ушбу шартномага мувофиқ қарздорликни тўлаш бўйича мажбуриятлари унинг кафиллари томонидан амалга оширилади, агарда унинг кафиллари томонидан хам қарздорликларни ундириш имконияти бўлмаса, Инвестторни амалда киритган инвестиция маблағлари Жамият томонидан пул ўтказиш йўли билан 50% (фоиз) миқдорида қайтарилган ҳолда шартнома бир томонланма суддан ташқари бекор қилинади.  
        3.2. Жамиятнинг ҳуқуқ ва мажбуриятлари:
        3.2.1. Жамият, ушбу шартноманинг предметига боғлиқ равишда Инвесторнинг инвестицион бадалини мақсадли сарф қилиши лозим.
        3.2.2. Жамият, лойиҳа-смета ҳужжатлари, шаҳарсозлик қоида ва меъёрларига мувофиқ равишда Турар жой биносининг қурилишини ва фойдаланишга топширилишини амалга ошириши лозим.
        3.2.3. Жамият, тавсифи ушбу шартноманинг 1.3. ва 1.5. бандларида келтирилган Объектни Инвесторга топшириш-қабул қилиш шартномасига мувофиқ, Инвестор томонидан инвестицион бадали киритиш  бўйича шартномавий мажбуриятлари тўлиқ бажарилганидан сўнг топшириши, шу жумладан, Инвесторга мулкка эгалик ҳуқуқини расмийлаштириш учун зарур бўлган барча ҳужжатларни тақдим қилиши лозим. 
        3.2.4. Жамият, объектив жиҳатдан шаҳарсозлик меъёр ва қоидаларнинг қўшимча талабларини бажариш билан боғлиқ бўлган ҳолларда, бир томонлама равишда, Объект қурилиши муддатларини ўзгартиришга ҳақли. 
        3.2.5. Инвестор инвестицион бадални киритиш муддатларини қўпол равишда бузган бўлса, Жамият бир томонлама шартномани бекор қилиши ва шартнома бўйича киритилган инвестицион бадалнинг 50 % (фоизини) қайтариб бермасликка ҳақли.  
        3.2.6. Жамият Инвесторда қарздорлик мавжуд бўлса, унинг ҳисоб рақамидан акцептсиз равишда пул маблағларини ечиб олиши мумкин.
"""
aaa="""
     4. Инвестициялар ҳажми ва ҳисоб-китоблар тартиби
"""
ab= """
        4.1. Инвестор томонидан киритилиши лозим бўлган инвестициялар ҳажми миллий валютада   _________________  (_________________________________) сўм ҚҚС билан ташкил қилади. Инвестициялар миқдорини ҳисоблаш учун шартноманинг 1.3. бандида кўрсатилган турар жой хонадонининг умумий қиймати ва квартиранинг бир квадрат метрининг қиймати Турар жой биносининг қурилишида иштирок этиш учун бадалнинг ушбу шартномага (№1-илова)га мувофиқ Инвестор ҳар ойлик тўловни ушбу шартнома имзоланганидан бошлаб ҳар ойнинг 15-санасидан кечиктирмаган ҳолда амалга ошириши лозим.
            4.2. Инвестор томонидан инвестицион бадалнинг киритилиши Ўзбекистон Республикасининг миллий валютасида Жамиятнинг ҳисоб рақамига амалга оширилади. 
        4.3. Инвестор томонидан инвестиция бадали пул маблағи шаклида киритилганлини тасдиқловчи ҳужжат Жамият ҳисоб рақамига пул маблағлари қабул қилинганлиги ҳақидаги банк белгиси қўйилган юқоридаги тўлов топшириқномасидан (квитанциясидан) иборатдир.
        4.4. Инвесторнинг инвестиция бадалини киритиш бўйича мажбуриятини бажарганлик санаси сифатида Жамият ҳисоб рақамига пул маблағлари келиб тушганлиги ҳақидаги Банк кўчирма ҳужжатида кўрсатилган сана ёки стандарт қурилиш материалларини топшириш-қабул қилиш далолатномасининг санаси қабул қилинади.
        4.5. Инвестор томонидан инвестиция бадали қурилиш материаллари шаклида киритилганлигини тасдиқловчи ҳужжат сифатида, томонлар имзолаган топшириш-қабул қилиш далолатномаси хизмат қилади.
        4.6. Инвестор Жамиятга инвестицион бадал тўловига эгалик қилиши, фойдаланиши, тасарруф қилиш бўйича чекланмаган ваколат бериб, ушбу шартномада кўрсатилган мақсадлар турар-жой биносини, шу жумладан, инвестиция Объектини қуриш ишларига йўналтиради.
        4.7. Шартноманинг 4.1. бандида кўрсатилган инвестициялар ҳажми вазиятлардан келиб чиққан ҳолда Инвестор хабардордорлигида ўзгариши мумкин. 
        4.8. Инвестор қуйидаги ҳуқуқларга эга: 
        - инвестиция суммасини олдиндан киритиш. Бунда тарафлар томонидан тасдиқланган бирламчи тўловлар жадвали суммаси, Жадвалга мувофиқ охирги ойлар учун камаяди (сўндирилади). 
        4.9. Агар Ўзбекистон республикасида энг кам ойлик иш ҳақи миқдори оширлиганда, шартноманинг (1-илова)сидаги тўлов жадвалида белгиланган суммалар инвестор билан келишилган холда қайта кўриб чиқилиши мумкин.
"""
aab= """
5. Алоҳида шартлар ва сифат кафолати
"""
aabb = """
        5.1. Инвесторнинг инвестиция Объектига нисбатан эгалик ҳуқуқи қонунчиликда ўрнатилган тартибда давлат рўйхатидан ўтказилган вақтдан бошлаб вужудга келади. 
        5.2. Инвестор томонидан олинаётган квартиранинг умумий майдони қурилиш якунига етганидан сўнг, бинонинг инвентаризацияси натижаларига кўра ва кадастр ҳужжатларига мувофиқ аниқланиши лозим.
        5.3. Тарафлар, объектнинг сифат кафолати, унинг лойиҳага, шаҳарсозлик меъёр ва қоидаларга мувофиқлигининг кўрсаткичи сифатида жойлардаги давлат хокимияти органи томонидан ўрнатилган тартибда тасдиқланган Турар жой биносини фойлаланишга қабул қилиш бўйича Комиссиянинг далолатномаси хизмат қилади.
        5.4. Инвестор, эгалик ҳуқуқи расмийлаштирилганидан сўнг ва инвестиция Объектини топшириш-қабул қилиш далолатномаси имзоланган санадан 10(ўн) кун давомида кўринмас кемтикликлар юзасидан Жамиятга эътироз билдириши мумкин. Барча эътирозлар асосланган ва ёзма равишда тузилган бўлиши лозим.
        5.5. Жамият эътирознома қабул қилинганидан сўнг 1 (бир)ойдавомида Инвесторнинг эътирози асосланганлигини кўриб чиқиши лозим ҳамда у асосланган бўлса, ўз маблағи ёки кучи ҳисобидан аниқланган кемтикликларни бартараф этади.
        5.6. Объектни топшириш муддати Тарафлар келишувига асосан узайтирилиши мумкин.
        Топшириш муддати, Тарафлар ушбу шартномани тузиш вақтида кўра билишининг иложи бўлмаган салмоқли вазият ўзгаришлари вужудга келганда, Объектни топшириш муддати узайтирилиши мумкин.
        5.7. Жамият томонидан турар-жой биносининг қавати ҳамда уй хонадон рақами объектив ва субъектив жараёнларга қараб пастки ёки устки қаватга ўзгариши мумкин. 
        Бундай вазиятда пастки қаватга ўзгарганда, Инвестор томонидан Жамиятга ўртадаги фарқ сумма тўлаб берилади, юқори қаватга ўзгарганда эса Жамият томонидан 1-иловага мувофиқ тузилган шартнома асосида тўланиши лозим  бўлган суммадан айирилади. 
"""
aba="""
 6. Тарафларнинг масъулияти ва зиддиятларни бартараф этиш тартиби
"""
aaba="""
        6.1. Инвестицион бадални киритиш муддатини бузган Инвестор Жамиятга ҳар бир кечиктирилган кун учун ушбу шартноманинг тўланмаган суммасининг 0,4 % миқдорида тўловни амалга оширади, бироқ бу сумма ушбу шартноманинг умумий қийматининг 50 % фоизидан ошмаслиги лозим. 
        6.2. Инвестор муттасил равишда ушбу шартномага мувофиқ тўлов киритиш мажбуриятини бузиб, икки ой ичида икки маротабадан ортиқ кечиктирса ёки тўлов бир ойдан ортиқ кечиктирилса, Жамият Инвесторни 5 (беш) календарь иш куни олдин ёзма равишда огоҳлантирган тарзда, суддан ташқари тартибда бир томонлама мазкур шартномани бекор қилиши мумкин. 
        Бу ҳолда Жамият бошқа Инвесторни жалб қилиб, Инвестор томонидан амалда киритилган инвестиция тўловларини Инвесторга шартноманинг бекор қилиниши ҳақидаги хабарнома жўнатилган кундан бошлаб 6 (олти) ой ичида 50 % (фоиз)миқдорида пул маблағини пул кўчириш йўли билан қайтаришга ҳақли. 
        6.3. Инвесторга Объектни эгаликка топшириш муддатларини асоссиз бузган Жамият Инвесторга ушбу шартноманинг умумий қийматидан 0.1 % миқдорида жарима тўлайди, бироқ бу сумма шартнома қийматининг 10 % фоизидан ошмаслиги лозим.
        6.4. Агар Жамиятнинг айби бўлмай туриб, Инвестор ушбу шартнома бўйича мажбуриятларини ижро қилишдан бир тарафлама бош тортса ва ўз ташаббусига кўра шартномани бекор қилса, Инвестор Жамиятга шартноманинг умумий қийматидан 10 % (ўн фоиз) миқдорида жарима тўлайди. Бу ҳолда Жамият ушбу шартнома бекор қилинганидан кундан бошлаб 60 (олтмиш)  банк куни ичида киритилган инвестиция суммасини жарима суммаси ва ушбу бандда кўрсатилган бошқа ҳаражатларни ажратган ҳолда, киритилган маблағларни қайтариб беради.
        6.5. Жамият томонидан ўз ушбу шартномада кўрсатилган мажбуриятларини бажариш бинонинг шаҳар муҳандислик тармоқларига уланиш ёки уларни мувофиқ идораларга топшириш билан боғлиқ равишда, учинчи шахсларнинг айби билан, кечиктирилса, Жамият ушбу шартномада кўрсатилган жавобгарликни ўз зиммасига олмайди.
        6.6. Инвестор ушбу шартномани ижро этишдан бир маротаба бош тортса ва Объект тўлови учун маблағ ўтказмаса, Жамият бир тарафлама тартибда Солиштирма далолатнома асосида Инвесторнинг банкига тўлов талабномасини чиқаришга ҳақли.
        6.7. Қолган қисмида, Тарафларнинг ушбу шартномага мувофиқ мажбуриятларини бажармаслик ёки керакли тарзда бажармаслик учун жавобгарлиги, Ўзбекистон Республикасининг Фуқаролик кодекси ва Ўзбекистон Республикасида амалда бўлган бошқа қонунчилик меъёрлари билан белгиланади.
        6.8. Ушбу шартномадан ёки у билан боғлиқ равишда келиб чиққан низо ва келишмовчиликлар Ўзбекистон Республикаси қонунчилиги билан ўрнатилган тартибда ҳал этилади.
"""
aaaba="7. Форс-мажор вазиятлар"
aaaaba="""
        7.1. Тарафларнинг ушбу шартномага мувофиқ мажбуриятларини бажаришларига тўсқинлик қилувчи енгиб бўлмас куч билан боғлиқ вазиятлар, яъни қурилиш тақиқланган сув тошқинлари, ёнғинлар, зилзила, об-ҳаво шароити, шунингдек, эпидемиялар, ҳарбий зиддиятлар, террорчилик ҳаракатлари, фуқаролик норозиликлари, норозилик намойишлари, давлат ҳокимияти органлари томонидан чиқарилган буйруқ, қарор, 
        фармойишлар ва бошқа чекловлар вужудга келса, ушбу шартномани ижро қилиш муддати ушбу ҳолатлар бартараф этилгунига қадар сурилади, ёки унинг форс-мажор вазиятлари билан чамбарчас боғлиқ қисми ижро этилмай туради.
        7.2. Агар, 7.1. бандда кўзда тутилган вазият вужудга келса, Тарафлардан бири иккинчисини бундай вазият вужудга келганлиги ҳақида 10 (ўн) календарь куни ичида 
        хабардор қилиши шарт. Бундай хабарномани жўнатмаган Тараф,  форс-мажор вазиятларни важ қилиб кўрсатиш ҳуқуқидан махрум бўлади.
        7.3. Агар форс-мажор вазиятлар узлуксиз ёки жамлашган ҳолда 2 (икки) ойдан ортиқ давом этса, ушбу Шартнома Тарафлардан ҳар бирининг ташаббуси билан бекор қилиниши лозим, бунда Жамият 60 (олтмиш)  банк куни ичида киритилган инвестиция суммасини Инвесторга қайтариб бериши лозим."""
abba="""
8. Шартноманинг ўзгартирилиши ва бекор қилиниши
"""
abbba="""
        8.1. Ўзбекистон Республикасининг қонунчилигида, ушбу шартнома шартларидан бири қонунчиликка мувофиқ келмай қолса, шартноманинг бошқа моддалари ва бандларининг ижро этилишига таъсир қилмайди. Қонунчиликка мувофиқ келмай қолган бандлар Тарафларнинг ўзаро келишувига мувофиқ маъно ва мақсад жиҳатидан энг яқин ва қонуний бўлган қоидага ўзгартирилади.
        8.2. Тарафлар томонидан шартнома қуйидаги ҳолларда муддатдан олдин бекор қилиниши мумкин:
        8.2.1. Тарафларнинг ўзаро келишувига кўра;
        8.2.2. Жамият ташаббуси билан;
        8.2.3. Инвесторнинг ташаббусига кўра;
        8.2.4. Форс-мажор вазиятлар туфайли.
        8.2.5. Шартноманинг муддатидан аввал бекор қилиниши ёзма равишда, Тарафлар бекор қилиш келишувини имзолаш йўли билан амалга оширилади. Шартноманинг бир тарафлама суддан ташқари бекор қилиниши бундан мустасно.
        8.2.6. Жамият томонидан вазиятдан келиб чиқиб шартнома бандларига Инвесторни хабардор қилган ҳолда ўзгартириш ва қўшимчалар киритилиши мумкин
        8.2.7 Тарафлардан бири ушбу шартномани ижро этишдан бир томонлама бош тортса, Шартнома иккинчи Тарафга шартноманинг бекор қилиниши ҳақида хабарномани жўнатган кунидан бошлаб кучга киради. Ушбу хабарнома почта орқали, буюртма хати билан иловалар рўйхати билан жўнатилиши лозим.
        8.2.8 Ушбу шартноманинг ҳар қандай ўзгартириш ва қўшимчалари ёзма равишда, ушбу Шартномага қўшимча келишувни расмийлаштириш йўли билан амалга оширилади. 
        8.9 Тарафлар ушбу шартномада кўзда тутилган ҳуқуқ ва мажбуриятларини учинчи шахсларга иккинчи Тарафнинг ёзма розилигисиз топшириш ҳуқуқига эга эмаслар.
        8.10. Ушбу шартномада кўзда тутилмаган масалаларда, Тарафлар Ўзбекистон Республикасида амалда бўлган қонунчилигига кўра иш олиб борадилар.
        8.11. Ушбу шартноманинг барча иловалари унинг ажралмас қисмини ташкил этади.
        8.12. Маълумотлари ва жойлашган манзили ўзгарган Тарафлар 5 (беш) кун ичида бир-бирларини рўй берган ўзгаришлар ҳақида хабрдор қилиши лозим. """
ababa="""
9. Якуний қоидалар
"""
hehe="""
        9.1. Ушбу шартнома имзоланган вақтидан бошлаб кучга киради ва Тарафлар ўз мажбуриятларини тўлиқ бажаришига ва ўзаро ҳисоб-китоблар якунланишига қадар амал қилади.
        9.2. Ушбу шартнома бир хил юридик кучга эга иккита асл нусхада тузилган бўлиб, Тарафларнинг ҳар бирига бир нусхадан тақдим этилади
        9.3. Ушбу шартнома Инвестор ҳисобидан нотариал тасдиқдан ўтиши лозим.
        9.4.  Шартноманинг барча нусхалари тенг юридик кучга эгадир.
"""
hehe1="""
10. Тарафларнинг манзиллари, реквизитлари ва имзолари
"""

document = Document()
style = document.styles['Normal']
p = document.add_paragraph()
p1 = document.add_paragraph()
p2 = document.add_paragraph()
p3 = document.add_paragraph()
p4 = document.add_paragraph()
p5 = document.add_paragraph()
p6 = document.add_paragraph()
p7 = document.add_paragraph()
p8 = document.add_paragraph()
p9 = document.add_paragraph()
p10 = document.add_paragraph()
p11 = document.add_paragraph()
p12 = document.add_paragraph()
p13 = document.add_paragraph()
p14 = document.add_paragraph()
p15 = document.add_paragraph()
p16 = document.add_paragraph()
p17 = document.add_paragraph()
p18 = document.add_paragraph()
p19 = document.add_paragraph()
p20 = document.add_paragraph()
p21 = document.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
p5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p6.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
p7.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p8.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
p9.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p10.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
p11.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p12.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
p13.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p14.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
p15.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p16.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
p17.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p18.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
p19.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p20.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
p21.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = p.add_run(x)
r2 = p1.add_run(x1)
r3= p2.add_run(y)
r4=p3.add_run(xx)
r44=p4.add_run(yy)
r5 =p5.add_run(yyy)
r6 = p6.add_run(yyyy)
r7 = p7.add_run(xxx)
r8 = p8.add_run(a)
r9 =p9.add_run(aaa)
r10 =p10.add_run(ab)
r11 = p11.add_run(aab)
r12 = p12.add_run(aabb)
r13 =p13.add_run(aba)
r14 = p14.add_run(aaba)
r15 = p15.add_run(aaaba)
r16 = p16.add_run(aaaaba)
r17 = p17.add_run(abba)
r18 = p18.add_run(abbba)
r19 =p19.add_run(ababa)
r20=p20.add_run(hehe)
r21=p21.add_run(hehe1)

r1.font.size = 16
r1.bold = True
r2.font.size = 12
r2.bold = True
r3.font.size = 12
r4.bold=True
r4.font.size=12
r44.font.size=8
r5.font.size = 12
r5.bold = True
r6.font.size = 12
r7.font.size = 12
r7.bold = True
r8.font.size = 12
r9.font.size = 12
r9.bold = True
r10.font.size=12
r11.font.size=12
r11.bold = True
r12.font.size=12
r13.font.size=12
r13.bold = True
r14.font.size=12
r15.font.size=12
r15.bold = True
r16.font.size=12
r17.font.size=12
r17.bold = True
r18.font.size=12
r19.font.size=12
r19.bold=True
r20.font.size=12
r21.font.size=12
r21.bold=True

data_of="""
Оқилов кўчаси, 18-уй, 28-хонадон
Х/Р: 2020 8000 700 749 187 002
“Халк банк” АК Тошкент шахар  ф-ли
МФО:01132, ИНН:304757296
ОКЭД: 61133  Тел: (95) 146-28-38
МЧЖ  директори

М.А.Очилов ________________
"""
data_of_1="""
Яшаш манзили: 


(СТИР: ______________, ЖШШИР: _______________) __ _________ ____________________________  томонидан ______________ йил берилган
Тел: (___________________________)  



_____________  ___________________
"""

data = (
    ("“INSOFT INVEST” МЧЖ", ' '),
    (data_of, data_of_1)
)
table = document.add_table(rows=1, cols=2)  
table.alignment = WD_TAB_ALIGNMENT.CENTER
row = table.rows[0].cells
row[0].text = "ЖАМИЯТ"
row[1].text = 'ИНВЕСТОР'
data_of_2 = row[0].paragraphs[0]
data_of_2.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
data_of_3 = row[1].paragraphs[0]
data_of_3.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
data_of_={}
for id, name in data:
    row = table.add_row().cells
    row[0].text = str(id)
    row[1].text = name
    data_of_[0] = row[0].paragraphs[0]
    data_of_[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    data_of_[1] = row[1].paragraphs[0]
    data_of_[1].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    

section = document.sections[0]
footer = section.footer
footer_para = footer.paragraphs[0]
footer.bold=True 
footer_para.text = "Жамият_____________\nИнвестор___________"

document.add_page_break()

document.save('demo.docx')
